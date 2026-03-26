import os
import io
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
import google.generativeai as genai
import streamlit as st
import pandas as pd
from sqlalchemy import create_engine, text
from datetime import datetime


load_dotenv()

genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
model = genai.GenerativeModel('gemini-2.5-pro')


def read_contituion(file_path):
    try:
        with open(file_path, 'r' , encoding='utf-8') as contituion:
            return contituion.read()
    except FileNotFoundError:
        return "Error: Constitution file not found!"
    
presidental_contituon = read_contituion('presidental_contituion.txt')



st.set_page_config(page_title="Job Searching App -with AI", layout="wide")
st.title("AI Job Finder")

st.sidebar.header("Searching Filters")

choice_city = st.sidebar.selectbox("PLease Choice City", ["Amsterdam", "Rotterdam", "Utrecht", "Eindhoven", "Helmond", "Den Haag"])

distance_km = st.sidebar.slider("Maximum Distance to the Center (km)", 0, 100, 25)

min_point = st.sidebar.slider("Minimum Job Score", 0, 100, 25)



my_profile= {
    "Name" : os.getenv("MY_NAME"),
    "Age" : os.getenv("MY_AGE"),
    "Location" : os.getenv("MY_LOCATION"),
    "Engineering Background" : os.getenv("MY_BACKGROUND"),
    "Target and Vision" : os.getenv("MY_TARGET_AND_VISION")
}

letter_rules = os.getenv("COVER_LETTER_RULES")
my_adress = os.getenv("MY_ADRESS", "").replace('\\n' , '\n')
cover_letter_hobby = os.getenv("COVER_LETTER_HOBBY", "").replace('\\n' , '\n')
cover_letter_closing = os.getenv("COVER_LETTER_CLOSING", "").replace('\\n' , '\n')
ai_override = os.getenv("AI_OVERRIDE", "").replace('\\n', '\n')
ai_format = os.getenv("AI_FORMAT", "").replace('\\n', '\n')
base_cv_info = os.getenv("BASE_CV_INFO", "").replace('\\n', '\n')
ai_cv_prompt = os.getenv("AI_CV_PROMPT", "").replace('\\n', '\n')
interview_secret_strategy = os.getenv("INTERVIEW_SECRET_STRATEGY", "").replace('\\n', '\n')


engine = create_engine('sqlite:///jobs.db')

with engine.connect() as conn:
    conn.execute(text("""
    CREATE TABLE IF NOT EXISTS applied_jobs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        company TEXT,
        job_title TEXT,
        applied_date TEXT
    )
    """))
    conn.commit()

df = pd.read_sql("SELECT * FROM job_posting", engine)

filtered_df = df[(df['location'].str.contains(choice_city, case=False, na=False))  & (df['ai_score'].astype(float) >= min_point)]

tab1, tab2 = st.tabs(["Job Search", "My Applications"])

with tab1:
    st.subheader(f"{choice_city} Jobs in the Surrounding Area")

    clean_df = filtered_df.drop(columns=['id', 'external_id'], errors='ignore')

    def points_colors(value):
        try:
            point = float(value)
        except:
            return ''
        
        if point >= 75:
            return 'background-color: #d9ead3; color: black'
        elif point >= 50:
            return 'background-color: #fff2cc; color: black'
        else:
            return 'background-color: #f4cccc; color: black'

    colorful_table = clean_df.style.map(points_colors, subset=['ai_score'])

    st.dataframe(colorful_table, hide_index=True)

    st.divider()
    st.subheader("🗣️ Talk to the Board Chairman (AI Chat)")

    chosen_advert = st.selectbox("Which listing do you want to process?", filtered_df['title'])

    if chosen_advert:
        st.write(f"**{chosen_advert}** Guide AI is now connected. You can ask your questions.")

        interview_mode = st.toggle("👔 Switch to Lead Interviewer Mode!")

        if "messages" not in st.session_state:
            st.session_state.messages = []

        for msg in st.session_state.messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
        
        if user_message := st.chat_input("Ask the guide or answer the interviewer..."):

            with st.chat_message("user"):
                st.markdown(user_message)
            st.session_state.messages.append({"role": "user", "content": user_message})

            chosen_job_detail = filtered_df[filtered_df['title'] == chosen_advert]['description'].values[0]
            job_page = filtered_df[filtered_df['title'] == chosen_advert]['company'].values[0]

            chat_history = "\n".join([f"{msg['role']}: {msg['content']}" for msg in st.session_state.messages[:-1]])

            if interview_mode:
                prompt_to_send = f"""
                Act as an insider Hiring Manager / Senior Tech Lead at {job_page} for the '{chosen_advert}' position.
                We are having an informal, pre-interview chat. You have deep knowledge of {job_page}'s real corporate culture, engineering standards, and tech stack.
                
                Candidate's CV Info: {base_cv_info}
                Job Description: {chosen_job_detail}
                
                Previous Conversation History:
                {chat_history}
                
                Candidate's Latest Message: {user_message}
                
                {interview_secret_strategy}
                
                YOUR INSTRUCTIONS:
                1. Analyze the candidate's last message. Provide direct, senior-level feedback. Tell them what was good, what was missing, and how to say it better in the actual interview.
                2. Share a "Golden Nugget" of insider information about {job_page}.
                3. Ask EXACTLY ONE highly specific, challenging technical or cultural question. Tailor it to the job description AND the candidate's transition from an Industrial/OT background to Software Engineering.
                4. Be pragmatic and realistic, but remember you are ultimately mentoring them to succeed.
                """
            else:
                prompt_to_send = f"""
                {presidental_contituon}

                Act as a friendly Career Guide and Assistant. 
                The user is applying for '{chosen_advert}' at {job_page}.
                Job Description: {chosen_job_detail}
                
                Previous Conversation History:
                {chat_history}
                
                User's Latest Question: {user_message}
                
                Answer the user's question clearly and concisely. Help them understand the job role, the company, or give general advice. Be supportive and motivating.
                """

            ai_response = model.generate_content(prompt_to_send).text

            with st.chat_message("assistant"):
                st.markdown(ai_response)
            st.session_state.messages.append({"role": "assistant", "content": ai_response})

        if st.button(f"{chosen_advert} for prepare a cover letter"):
            today_date = datetime.now().strftime("%B %d, %Y")
            job_page = filtered_df[filtered_df['title'] == chosen_advert]['company'].values[0]
            chosen_job_detail = filtered_df[filtered_df['title'] == chosen_advert]['description'].values[0]

            ai_instructions = f"""
            {ai_override}

            {presidental_contituon}

            Job Detail: {chosen_job_detail}
            Company: {job_page}

            {letter_rules}

            {ai_format}
            """

            ai_motivation_letter = model.generate_content(ai_instructions).text.strip()

            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(11)

            adress_para = doc.add_paragraph()
            adress_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            adress_run = adress_para.add_run(my_adress)
            adress_run.font.size = Pt(12)

            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"\n{today_date}, Helmond\n\n")
            date_run.font.size = Pt(12)

            doc.add_paragraph("Dear Hiring Manager,\n")
            doc.add_paragraph(ai_motivation_letter)
            doc.add_paragraph(cover_letter_hobby + "\n")
            doc.add_paragraph(cover_letter_closing)

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.success("Motivation Letter is ready!")

            st.download_button(
                label="⬇️ Download Motivation Letter",
                data=buffer,
                file_name=f"Ozan_Motivation_Letter_{job_page}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        if st.button(f"{chosen_advert} for Generate Portfolio Summary"):
            chosen_job_detail = filtered_df[filtered_df['title'] == chosen_advert]['description'].values[0]
            job_page = filtered_df[filtered_df['title'] == chosen_advert]['company'].values[0]

            ai_cv_instructions = f"""
            {ai_cv_prompt}

            {base_cv_info}

            {chosen_job_detail}
            """

            ai_cv_content = model.generate_content(ai_cv_instructions).text.strip()

            try:
                summary_text = ai_cv_content.split("[TECHNICAL SKILLS]")[0].replace("[CV SUMMARY]", "").strip()
                skills_text = ai_cv_content.split("[TECHNICAL SKILLS]")[1].strip()
            except Exception as e:
                summary_text = "Parsing error. AI output format was unexpected"
                skills_text = ai_cv_content

            doc = DocxTemplate("cv_template.docx")

            context = {
                'SUMMARY' : summary_text,
                'SKILLS' : skills_text
            }

            doc.render(context)

            cv_buffer = io.BytesIO()
            doc.save(cv_buffer)
            cv_buffer.seek(0)

            st.success("Automated CV is ready with your custom format!")

            st.download_button(
                label="⬇️ Adapted CV Download",
                data=cv_buffer,
                file_name=f"Ozan_CV_{chosen_advert}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        st.divider()
        if st.button(f"Mark as Applied for {chosen_advert}"):
            apply_date = datetime.now().strftime("%Y-%m-%d")
            job_page = filtered_df[filtered_df['title'] == chosen_advert]['company'].values[0]

            with engine.connect() as conn:
                conn.execute(text(f"INSERT INTO applied_jobs (company, job_title, applied_date) VALUES ('{job_page}', '{chosen_advert}', '{apply_date}')"))
                conn.commit()

            st.success("Successfully marked as applied! Check 'My Applications' tab.")

if not filtered_df.empty:
    st.write("### 🧠 AI Analysis for Best Matches")

    top_job = filtered_df.sort_values(by='ai_score', ascending=False).iloc[0]

    with st.expander(f"Why did AI give {top_job['ai_score']} points to {top_job['title']}?"):
        st.info(top_job['ai_reasoning'])

with tab2:
    st.subheader("My Job Applications")
    try:
        applied_jobs_df = pd.read_sql("SELECT * FROM applied_jobs", engine)
        if not applied_jobs_df.empty:
            clean_applied_df = applied_jobs_df.drop(columns=['id'], errors='ignore')
            st.dataframe(clean_applied_df, hide_index=True, use_container_width=True)
        else:
            st.info("You haven't applied for any jobs yet. Let's find one!")
    except Exception as e:
        st.info("No Application data found")